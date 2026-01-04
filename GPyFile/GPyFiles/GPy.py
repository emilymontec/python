import os
import sys
import shutil
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
import json

# Obtener la ruta del directorio donde se encuentra el script
if getattr(sys, 'frozen', False):
    SCRIPT_DIR = os.path.dirname(sys.executable)
else:
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
GPYFILES_DIR = os.path.join(SCRIPT_DIR, ".gpyfiles")
ARCHIVOS_DIR = os.path.join(GPYFILES_DIR, "archivos")
PAPELERA_DIR = os.path.join(GPYFILES_DIR, "papelera")

def resolve_path(path):
    if os.path.isabs(path):
        return path
    else:
        return os.path.join(ARCHIVOS_DIR, path)

def listar(ruta=ARCHIVOS_DIR):
    print(f"\nContenido de {os.path.abspath(ruta)}:")
    for elemento in os.listdir(ruta):
        tipo = "[DIR]" if os.path.isdir(os.path.join(ruta, elemento)) else "[FILE]"
        print(f"{tipo} {elemento}")
    print()

def crear_archivo(nombre):
    # Si no se especifica una ruta, se crea en la carpeta 'archivos'
    if os.path.dirname(nombre) == '':
        nombre = os.path.join(ARCHIVOS_DIR, nombre)

    extension = os.path.splitext(nombre)[1].lower()

    # Crea según la extensión
    if extension in [".txt", ""]:
        with open(nombre, "w", encoding="utf-8") as f:
            f.write("")
    elif extension == ".docx":
        doc = Document()
        doc.add_paragraph("Documento creado automáticamente.")
        doc.save(nombre)
    elif extension == ".pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="PDF generado automáticamente.", ln=True)
        pdf.output(nombre)
    elif extension in [".xlsx", ".xls"]:
        wb = Workbook()
        ws = wb.active
        ws.title = "Hoja1"
        ws["A1"] = "Archivo Excel generado automáticamente."
        wb.save(nombre)
    elif extension == ".csv":
        with open(nombre, "w", encoding="utf-8") as f:
            f.write("columna1,columna2\nvalor1,valor2\n")
    elif extension == ".json":
        with open(nombre, "w", encoding="utf-8") as f:
            json.dump({"mensaje": "Archivo JSON creado automáticamente."}, f, indent=4)
    else:
        with open(nombre, "w", encoding="utf-8") as f:
            f.write("")  # Archivo genérico vacío
    print(f"Archivo '{nombre}' creado.")
    return nombre # Devolvemos la ruta completa

def crear_carpeta(nombre):
    # Si no se especifica una ruta, se crea en la carpeta 'archivos'
    if os.path.dirname(nombre) == '':
        nombre = os.path.join(ARCHIVOS_DIR, nombre)
    os.makedirs(nombre, exist_ok=True)
    print(f"Carpeta '{nombre}' creada.")

def eliminar(ruta):
    ruta = resolve_path(ruta)
    if not os.path.exists(ruta):
        print("Ruta no encontrada.")
        return

    # Si el archivo ya está en la papelera, pide confirmación para borrarlo permanentemente
    if os.path.abspath(ruta).startswith(PAPELERA_DIR):
        confirmacion = input(f"¿Seguro que quieres eliminar '{ruta}' permanentemente? (s/n): ").lower()
        if confirmacion == 's':
            if os.path.isdir(ruta):
                shutil.rmtree(ruta)
            else:
                os.remove(ruta)
            print(f"'{ruta}' eliminado permanentemente.")
        else:
            print("Eliminación cancelada.")
    else:
        # Mueve el archivo o carpeta a la papelera
        shutil.move(ruta, os.path.join(PAPELERA_DIR, os.path.basename(ruta)))
        print(f"'{ruta}' movido a la papelera.")

def renombrar(ruta, nuevo_nombre):
    ruta = resolve_path(ruta)
    if not os.path.exists(ruta):
        print("El archivo o carpeta no existe.")
        return None
    # Si el nuevo nombre no tiene ruta, se asume que es en el mismo directorio
    if os.path.dirname(nuevo_nombre) == '':
        nuevo_nombre = os.path.join(os.path.dirname(ruta), nuevo_nombre)
    else:
        nuevo_nombre = resolve_path(nuevo_nombre)
    os.rename(ruta, nuevo_nombre)
    print(f"Renombrado '{ruta}' a '{nuevo_nombre}'.")
    return nuevo_nombre # Devolvemos la nueva ruta completa

def mover(origen, destino):
    origen = resolve_path(origen)
    if not os.path.exists(origen):
        print("El archivo o carpeta no existe.")
        return
    if not os.path.isabs(destino):
        destino = os.path.join(ARCHIVOS_DIR, destino)
    shutil.move(origen, destino)
    print(f"'{origen}' movido a '{destino}'.")

def put_contenido(nombre_archivo, contenido):
    nombre_archivo = resolve_path(nombre_archivo)
    if not os.path.exists(nombre_archivo):
        print(f"El archivo '{nombre_archivo}' no existe. Créalo primero.")
        return

    extension = os.path.splitext(nombre_archivo)[1].lower()

    if extension == ".docx":
        try:
            doc = Document(nombre_archivo)
        except Exception:
            doc = Document()
        doc.add_paragraph(contenido)
        doc.save(nombre_archivo)
    elif extension in [".xlsx", ".xls"]:
        try:
            wb = Workbook()
            if os.path.exists(nombre_archivo):
                from openpyxl import load_workbook
                wb = load_workbook(nombre_archivo)
            ws = wb.active
            nueva_fila = ws.max_row + 1
            ws.cell(row=nueva_fila, column=1, value=contenido)
            wb.save(nombre_archivo)
        except Exception as e:
            print(f"Error al escribir en el archivo Excel: {e}")
            return
    elif extension == ".json":
        try:
            datos = {}
            if os.path.exists(nombre_archivo) and os.path.getsize(nombre_archivo) > 0:
                with open(nombre_archivo, "r", encoding="utf-8") as f:
                    datos = json.load(f)
            nueva_clave = f"entrada_{len(datos) + 1}"
            datos[nueva_clave] = contenido
            with open(nombre_archivo, "w", encoding="utf-8") as f:
                json.dump(datos, f, indent=4)
        except json.JSONDecodeError:
            print(f"Error: El archivo '{nombre_archivo}' no contiene un JSON válido. No se pudo agregar contenido.")
            return
    elif extension in [".txt", ".csv", ""]:
        with open(nombre_archivo, "a", encoding="utf-8") as f:
            f.write(contenido + "\n")
    else:
        print(f"La función 'put' no es compatible con archivos de extensión '{extension}'.")
        return
    print(f"Contenido agregado a '{nombre_archivo}'.")

def view_contenido(nombre_archivo):
    nombre_archivo = resolve_path(nombre_archivo)
    if not os.path.exists(nombre_archivo):
        print(f"El archivo '{nombre_archivo}' no existe.")
        return

    extension = os.path.splitext(nombre_archivo)[1].lower()
    print(f"\n--- Contenido de {nombre_archivo} ---\n")

    try:
        if extension in [".txt", ".csv", ""]:
            nombre_base = os.path.basename(nombre_archivo)
            with open(nombre_archivo, "r", encoding="utf-8") as f:
                for linea in f:
                    # Si la línea empieza con el nombre del archivo, lo quitamos
                    if linea.strip().startswith(nombre_base):
                        print(linea.strip()[len(nombre_base):].strip())
                    else:
                        print(linea, end="")
        elif extension == ".json":
            with open(nombre_archivo, "r", encoding="utf-8") as f:
                datos = json.load(f)
                print(json.dumps(datos, indent=4))
        elif extension == ".docx":
            doc = Document(nombre_archivo)
            for para in doc.paragraphs:
                print(para.text)
        elif extension in [".xlsx", ".xls"]:
            from openpyxl import load_workbook
            wb = load_workbook(nombre_archivo, data_only=True)
            ws = wb.active
            for row in ws.iter_rows():
                print(" | ".join([str(cell.value) if cell.value is not None else "" for cell in row]))
        else:
            print(f"La visualización para la extensión '{extension}' no es soportada en la consola.")
    except Exception as e:
        print(f"No se pudo leer el archivo '{nombre_archivo}': {e}")
    
    print(f"\n--- Fin del contenido ---")

def restaurar_archivo(nombre_archivo):
    ruta_papelera = os.path.join(PAPELERA_DIR, nombre_archivo)
    if not os.path.exists(ruta_papelera):
        print(f"El archivo '{nombre_archivo}' no se encuentra en la papelera.")
        return
    
    ruta_destino = os.path.join(ARCHIVOS_DIR, nombre_archivo)
    shutil.move(ruta_papelera, ruta_destino)
    print(f"'{nombre_archivo}' restaurado a la carpeta 'archivos'.")

def ayuda():
    print("""
Comandos disponibles:
  list [ruta]                -> Lista archivos y carpetas (por defecto en 'archivos').
  create <nombre.archivo>    -> Crea un archivo en la carpeta 'archivos'.
  mkdir <carpeta>            -> Crea una carpeta en 'archivos'.
  remove <ruta>              -> Mueve archivo/carpeta a la papelera.
                             -> Si está en la papelera, pide confirmación para borrar permanentemente.
  restore <nombre.archivo>   -> Restaura un archivo de la papelera a la carpeta 'archivos'.
  rename <viejo> <nuevo>     -> Renombra archivo o carpeta
  move <origen> <destino>     -> Mueve archivo o carpeta
  put "contenido"            -> Agrega contenido al último archivo creado/modificado.
  view <nombre.archivo>      -> Muestra el contenido de un archivo.
  help                       -> Muestra este mensaje
  exit                       -> Sale del programa
""")

# g:\Mi unidad\EVIDENCIAS\emc\GPyFiles\GPyFiles\GPy.py
# ... (otras importaciones y funciones) ...

if __name__ == "__main__":
    from GPyFilesCover import mostrar_portada, NEGRITA, RESET, AMARILLO_NEG
    mostrar_portada() # Aquí se llama a la función para mostrar la portada.
    print("\n" + NEGRITA + "....................Gestor de Archivos en Python...................." + RESET + "\n")
    print(NEGRITA + "......Soporta varios formatos: txt, docx, pdf, xlsx, csv, json......" + RESET + "\n")
    print(NEGRITA + "      Escribe 'help' para ver los comandos disponibles" + RESET + "\n")

    # Crear carpetas 'archivos' y 'papelera' si no existen
    os.makedirs(ARCHIVOS_DIR, exist_ok=True)
    os.makedirs(PAPELERA_DIR, exist_ok=True)

    ultimo_archivo = None

    while True:
        comando = input(AMARILLO_NEG + "> " + RESET).strip()
        if not comando:
            continue

        partes = comando.split()
        cmd = partes[0].lower()

        try:
            if cmd == "list":
                if len(partes) > 1:
                    ruta = partes[1]
                    if not os.path.isabs(ruta):
                        ruta = os.path.join(ARCHIVOS_DIR, ruta)
                else:
                    ruta = ARCHIVOS_DIR
                listar(ruta)
            elif cmd == "create" and len(partes) > 1:
                ultimo_archivo = crear_archivo(partes[1])
            elif cmd == "mkdir" and len(partes) > 1:
                crear_carpeta(partes[1])
            elif cmd == "remove" and len(partes) > 1:
                eliminar(partes[1])
            elif cmd == "rename" and len(partes) > 2:
                nueva_ruta = renombrar(partes[1], partes[2])
                if nueva_ruta:
                    ultimo_archivo = nueva_ruta
            elif cmd == "move" and len(partes) > 2:
                mover(partes[1], partes[2])
            elif cmd == "put" and len(partes) > 1:
                if ultimo_archivo:
                    contenido = " ".join(partes[1:])
                    if contenido.startswith('"') and contenido.endswith('"'):
                        contenido = contenido[1:-1]
                    put_contenido(ultimo_archivo, contenido)
                else:
                    print("Primero usa 'create' o 'rename' sobre un archivo para saber dónde agregar el contenido.")
            elif cmd == "view":
                if len(partes) > 1:
                    view_contenido(partes[1])
                elif ultimo_archivo:
                    view_contenido(ultimo_archivo)
                else:
                    print("Especifica un archivo para ver o usa 'create'/'rename' primero.")
            elif cmd == "restore" and len(partes) > 1:
                restaurar_archivo(partes[1])
            elif cmd == "help":
                ayuda()
            elif cmd == "exit":
                print("Saliendo del gestor de archivos...")
                break
            else:
                print("Comando no válido. Escribe 'help' para ver la lista.")
        except Exception as e:
            print(f"Error: {e}")
